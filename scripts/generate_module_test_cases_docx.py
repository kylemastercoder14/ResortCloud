from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "module-test-cases.md"
OUTPUT = ROOT / "docs" / "resortcloud-module-test-cases.docx"


def clean_markdown(value: str) -> str:
    value = value.strip()
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = value.replace("&", "&")
    return value


def parse_cases() -> list[dict[str, str]]:
    cases: list[dict[str, str]] = []
    section = ""
    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if line.startswith("## "):
            section = clean_markdown(line.removeprefix("## "))
            continue
        if not line.startswith("|"):
            continue
        cells = [clean_markdown(cell) for cell in line.strip("|").split("|")]
        if len(cells) < 4:
            continue
        if cells[0] in {"Module", "---"} or cells[0].startswith("---"):
            continue
        module, route, status, test_cases = cells[:4]
        if not route.startswith("/tenant") and route != "/accept-invitation":
            continue
        cases.append(
            {
                "section": section,
                "module": module,
                "route": route,
                "status": status,
                "test_cases": test_cases,
            }
        )
    return cases


def set_cell_border(cell, color: str = "000000", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(paragraph, text: str, size: int = 12) -> None:
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_plain_heading(doc: Document, text: str, size: int = 12):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(6)
    add_heading(paragraph, text, size)
    return paragraph


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(item)


def split_checks(text: str) -> list[str]:
    parts = [part.strip() for part in text.split(";") if part.strip()]
    if not parts:
        return ["Module behavior is verified."]
    return [part if part.endswith(".") else f"{part}." for part in parts]


def add_inputs_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.6)
    set_cell_text(table.rows[0].cells[0], "Input", bold=True)
    set_cell_text(table.rows[0].cells[1], "Value", bold=True)
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key)
        set_cell_text(cells[1], value)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)


def add_numbered_steps(doc: Document, steps: list[str]) -> None:
    for step in steps:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(step)


def build_docx() -> None:
    cases = parse_cases()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for style_name in ("List Bullet", "List Number"):
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.size = Pt(12)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(12)
    add_heading(title, "ResortCloud Module Test Cases", 14)

    intro = doc.add_paragraph()
    intro.add_run("Source: docs/module-test-cases.md")
    intro.paragraph_format.space_after = Pt(6)
    rules = doc.add_paragraph()
    rules.add_run(
        "Status rules: Passed means route exists with page.tsx. Coming soon means nav item is locked or route page does not exist yet."
    )
    rules.paragraph_format.space_after = Pt(12)

    for index, case in enumerate(cases, start=1):
        if index > 1:
            doc.add_paragraph()

        add_plain_heading(doc, f"Case {index} - {case['module']}", 12)

        add_plain_heading(doc, f"3.{index}.1 Purpose", 12)
        purpose = doc.add_paragraph()
        if case["status"] == "Passed":
            purpose.add_run(
                f"To verify that the {case['module']} module opens correctly, performs its expected workflow, and remains stable for authorized tenant users."
            )
        else:
            purpose.add_run(
                f"To verify that the {case['module']} module is correctly treated as coming soon, locked, or unavailable until implementation is complete."
            )

        add_plain_heading(doc, f"3.{index}.2 Inputs", 12)
        add_inputs_table(
            doc,
            [
                ("Module", case["module"]),
                ("Route", case["route"]),
                ("Section", case["section"]),
                ("Status", case["status"]),
            ],
        )

        add_plain_heading(doc, f"3.{index}.3 Expected Outputs & Pass/Fail Criteria", 12)
        doc.add_paragraph("Expected Output")
        expected = split_checks(case["test_cases"])
        if case["status"] != "Passed":
            expected = [
                "Coming soon or locked state is shown where applicable.",
                "Parent navigation expands without exposing an unfinished workflow.",
                "No broken production workflow is shown to the user.",
            ]
        add_bullets(doc, expected)

        doc.add_paragraph("Pass Criteria")
        if case["status"] == "Passed":
            add_bullets(
                doc,
                [
                    "Route opens without 404.",
                    "Module UI renders without runtime errors.",
                    "Listed module test cases pass.",
                ],
            )
        else:
            add_bullets(
                doc,
                [
                    "Coming soon or locked state is visible.",
                    "No unfinished route is presented as a completed feature.",
                ],
            )

        doc.add_paragraph("Fail Criteria")
        if case["status"] == "Passed":
            add_bullets(
                doc,
                [
                    "Route returns 404.",
                    "Page crashes or blocks the main workflow.",
                    "One or more listed test cases fail.",
                ],
            )
        else:
            add_bullets(
                doc,
                [
                    "Locked or missing module appears as fully available.",
                    "Navigation exposes a broken route or empty production page.",
                ],
            )

        add_plain_heading(doc, f"3.{index}.4 Test Procedure", 12)
        if case["status"] == "Passed":
            add_numbered_steps(
                doc,
                [
                    f"Open {case['route']}.",
                    "Verify that the page loads without a 404.",
                    "Run the expected module checks listed above.",
                    "Refresh the page and confirm the state remains stable.",
                    "Record the result as passed or failed.",
                ],
            )
        else:
            add_numbered_steps(
                doc,
                [
                    "Open the tenant sidebar or navigation menu.",
                    f"Locate {case['module']}.",
                    "Confirm the item is locked, marked coming soon, or not exposed as a finished workflow.",
                    "Confirm no broken page is reachable from the navigation.",
                    "Record the result as passed or failed.",
                ],
            )

        if index % 3 == 0 and index != len(cases):
            doc.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)

    props = doc.core_properties
    props.author = "ResortCloud"
    props.title = "ResortCloud Module Test Cases"
    props.subject = "Module test cases"
    props.keywords = "ResortCloud, test cases, QA"
    props.comments = "Generated from docs/module-test-cases.md"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)
    print(f"cases={len(cases)}")


if __name__ == "__main__":
    build_docx()
